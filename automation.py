from docx import Document

document = Document("doc.docx")

def extrair_matriculas(caminho_arquivo):

    document = Document("doc.docx")
    matriculas = []

    for paragrafo in document.paragraphs:
        texto = paragrafo.text
        if "==" in texto:
            matricula = texto.split("==")[1].strip()
            matriculas.append(matricula)
            document.save('teste', 1,'.docx')



    return matriculas


caminho_do_arquivo = "doc.docx"
matriculas_encontradas = extrair_matriculas(caminho_do_arquivo)


print(matriculas_encontradas)