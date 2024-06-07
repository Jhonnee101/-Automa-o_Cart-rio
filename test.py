import openpyxl as xl
import docx


workbook = xl.load_workbook("doc.xlsx")
planilha = workbook.active

for rows in planilha.iter_rows(min_row=2):
    nome, matricula = rows[0].value, rows[1].value

    # Crie um novo documento .docx
    doc = docx.Document()
    doc.add_paragraph(f"Nome: {nome}")
    doc.add_paragraph(f"Matrícula: {matricula}")

    nome_arquivo = f"{matricula}.docx"
    doc.save(nome_arquivo)


workbook.close()


